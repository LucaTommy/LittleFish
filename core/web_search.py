"""
Web search and research layer for Little Fish.

Uses Brave Search API for web queries and Groq for summarization.
All functions return plain strings suitable for TTS / Fish speech.
"""

import json
import os
import re
import traceback
from pathlib import Path
from typing import Optional

import requests
from bs4 import BeautifulSoup


_TIMEOUT = 10  # seconds for all web requests

_BRAVE_SEARCH_URL = "https://api.search.brave.com/res/v1/web/search"
_BRAVE_NEWS_URL = "https://api.search.brave.com/res/v1/news/search"
_BRAVE_IMAGE_URL = "https://api.search.brave.com/res/v1/images/search"

_NO_KEY_MSG = "Web search not configured. Add Brave API key in Settings."


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_brave_key() -> str:
    """Load Brave Search API key from secrets.json."""
    try:
        from config import load_secrets
        return load_secrets().get("brave_api_key", "")
    except Exception:
        return ""


def _get_groq_client(groq_keys: list):
    """Build a Groq client from the first available key."""
    if not groq_keys:
        return None
    try:
        import groq
        return groq.Groq(api_key=groq_keys[0])
    except Exception:
        return None


def _groq_complete(groq_keys: list, system: str, user: str,
                   max_tokens: int = 400, temperature: float = 0.3) -> Optional[str]:
    """Call Groq with key rotation on failure. Returns text or None."""
    if not groq_keys:
        return None
    try:
        import groq as groq_mod
    except ImportError:
        return None

    last_err = None
    for key in groq_keys:
        try:
            client = groq_mod.Groq(api_key=key)
            resp = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
                max_tokens=max_tokens,
                temperature=temperature,
            )
            return resp.choices[0].message.content.strip()
        except Exception as e:
            last_err = e
            continue

    print(f"[SEARCH] Groq failed after all keys: {last_err}")
    return None


def _brave_get(url: str, params: dict) -> Optional[dict]:
    """Make an authenticated request to Brave Search API."""
    key = _get_brave_key()
    if not key:
        return None
    headers = {
        "Accept": "application/json",
        "Accept-Encoding": "gzip",
        "X-Subscription-Token": key,
    }
    resp = requests.get(url, params=params, headers=headers, timeout=_TIMEOUT)
    resp.raise_for_status()
    return resp.json()


def _extract_page_text(html: str) -> str:
    """Extract readable text from HTML, stripping scripts/nav/footer."""
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header",
                     "aside", "form", "iframe", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    # Collapse whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


# ---------------------------------------------------------------------------
# Web Search
# ---------------------------------------------------------------------------

def web_search(query: str, groq_keys: list, max_results: int = 5) -> str:
    """Search the web via Brave API and summarize results with Groq."""
    try:
        if not _get_brave_key():
            return _NO_KEY_MSG

        print(f"[SEARCH] Web search: {query}")
        data = _brave_get(_BRAVE_SEARCH_URL, {"q": query, "count": max_results})
        if not data:
            return _NO_KEY_MSG

        results = data.get("web", {}).get("results", [])
        if not results:
            return f"No results found for '{query}'."

        # Build context for Groq
        context_parts = []
        sources = []
        for i, r in enumerate(results[:max_results], 1):
            title = r.get("title", "")
            url = r.get("url", "")
            desc = r.get("description", "")
            context_parts.append(f"{i}. {title}\n{desc}")
            sources.append(f"{title}: {url}")

        context = "\n\n".join(context_parts)

        # Summarize with Groq
        summary = _groq_complete(
            groq_keys,
            system="Answer the user's question based on these search results. "
                   "Be concise and conversational (2-4 sentences). "
                   "Speak naturally as if telling a friend.",
            user=f"Question: {query}\n\nSearch results:\n{context}",
            max_tokens=300,
        )

        if summary:
            source_list = ". ".join(sources[:3])
            return f"{summary}\n\nSources: {source_list}"

        # Fallback: just return raw results
        raw = ". ".join(f"{r.get('title', '')}: {r.get('description', '')}"
                        for r in results[:3])
        return f"Here's what I found: {raw}"

    except Exception as e:
        print(f"[SEARCH] web_search error: {e}")
        traceback.print_exc()
        return f"Search failed: {e}"


# ---------------------------------------------------------------------------
# Fetch & Summarize a Page
# ---------------------------------------------------------------------------

def web_fetch_page(url: str, groq_keys: list) -> str:
    """Fetch a webpage and summarize its content with Groq."""
    try:
        print(f"[SEARCH] Fetching page: {url}")
        resp = requests.get(
            url, timeout=_TIMEOUT,
            headers={"User-Agent": "LittleFish/1.5 (Desktop Companion)"},
        )
        resp.raise_for_status()

        text = _extract_page_text(resp.text)
        if not text.strip():
            return "Couldn't extract any readable content from that page."

        # Take first 3000 chars
        text = text[:3000]

        summary = _groq_complete(
            groq_keys,
            system="Summarize the key information from this webpage in 3-5 spoken sentences. "
                   "Be concise and conversational.",
            user=f"Webpage content:\n{text}",
            max_tokens=300,
        )

        return summary or f"Page content (first 500 chars): {text[:500]}..."

    except requests.exceptions.Timeout:
        return "The page took too long to load."
    except Exception as e:
        print(f"[SEARCH] web_fetch_page error: {e}")
        return f"Couldn't fetch that page: {e}"


# ---------------------------------------------------------------------------
# Research → Word Document
# ---------------------------------------------------------------------------

def web_research_to_doc(topic: str, groq_keys: list) -> str:
    """Research a topic using web search and create a Word document on Desktop."""
    try:
        if not _get_brave_key():
            return _NO_KEY_MSG

        print(f"[SEARCH] Researching topic: {topic}")

        # 1. Search for the topic
        data = _brave_get(_BRAVE_SEARCH_URL, {"q": topic, "count": 5})
        if not data:
            return _NO_KEY_MSG

        results = data.get("web", {}).get("results", [])
        if not results:
            return f"No results found for '{topic}'."

        # 2. Fetch content from top 3 URLs
        all_content = []
        sources = []
        for r in results[:3]:
            url = r.get("url", "")
            title = r.get("title", "")
            sources.append(f"{title} — {url}")
            try:
                page_resp = requests.get(
                    url, timeout=_TIMEOUT,
                    headers={"User-Agent": "LittleFish/1.5 (Desktop Companion)"},
                )
                page_text = _extract_page_text(page_resp.text)[:2000]
                all_content.append(f"Source: {title}\n{page_text}")
            except Exception:
                all_content.append(f"Source: {title}\n{r.get('description', '')}")

        combined = "\n\n---\n\n".join(all_content)

        # 3. Groq synthesizes into structured report
        report = _groq_complete(
            groq_keys,
            system=(
                "You are a research assistant. Write a structured report based on the provided sources. "
                "Format your response EXACTLY as:\n"
                "TITLE: [topic title]\n"
                "INTRODUCTION:\n[2 paragraphs]\n"
                "KEY FINDINGS:\n- [bullet point]\n- [bullet point]\n...\n"
                "DETAILS:\n[3-4 paragraphs]\n"
                "Do NOT include a sources section — that will be added automatically."
            ),
            user=f"Research topic: {topic}\n\nSources:\n{combined}",
            max_tokens=1500,
            temperature=0.4,
        )

        if not report:
            return "Couldn't generate the research report."

        # 4. Create Word document
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document()

        # Parse the report sections
        lines = report.split("\n")
        current_section = None
        section_content = []

        def _flush_section():
            nonlocal section_content
            if not current_section or not section_content:
                section_content = []
                return
            text = "\n".join(section_content).strip()
            if not text:
                section_content = []
                return

            if current_section == "TITLE":
                p = doc.add_heading(text, level=0)
                for run in p.runs:
                    run.font.size = Pt(18)
            elif current_section == "INTRODUCTION":
                for para in text.split("\n\n"):
                    p = doc.add_paragraph(para.strip())
                    for run in p.runs:
                        run.font.size = Pt(11)
            elif current_section == "KEY FINDINGS":
                doc.add_heading("Key Findings", level=1)
                for line in text.split("\n"):
                    line = line.strip().lstrip("- •")
                    if line:
                        p = doc.add_paragraph(line, style="List Bullet")
                        for run in p.runs:
                            run.font.size = Pt(11)
            elif current_section == "DETAILS":
                doc.add_heading("Details", level=1)
                for para in text.split("\n\n"):
                    p = doc.add_paragraph(para.strip())
                    for run in p.runs:
                        run.font.size = Pt(11)
            section_content = []

        for line in lines:
            stripped = line.strip()
            if stripped.startswith("TITLE:"):
                _flush_section()
                current_section = "TITLE"
                title_text = stripped[6:].strip()
                if title_text:
                    section_content.append(title_text)
            elif stripped == "INTRODUCTION:" or stripped.startswith("INTRODUCTION:"):
                _flush_section()
                current_section = "INTRODUCTION"
            elif stripped == "KEY FINDINGS:" or stripped.startswith("KEY FINDINGS:"):
                _flush_section()
                current_section = "KEY FINDINGS"
            elif stripped == "DETAILS:" or stripped.startswith("DETAILS:"):
                _flush_section()
                current_section = "DETAILS"
            else:
                section_content.append(line)

        _flush_section()

        # Sources section
        doc.add_heading("Sources", level=1)
        for src in sources:
            p = doc.add_paragraph(src)
            for run in p.runs:
                run.font.size = Pt(9)
                run.italic = True

        # 5. Save to Desktop
        desktop = Path.home() / "Desktop"
        safe_name = re.sub(r'[^\w\s-]', '', topic).strip().replace(" ", "_")[:50]
        filename = f"{safe_name}_research.docx"
        filepath = desktop / filename
        doc.save(str(filepath))

        print(f"[SEARCH] Research saved to {filepath}")
        return f"Research complete. Saved {filename} to your Desktop."

    except ImportError:
        return "python-docx package not available. Can't create documents."
    except Exception as e:
        print(f"[SEARCH] web_research_to_doc error: {e}")
        traceback.print_exc()
        return f"Research failed: {e}"


# ---------------------------------------------------------------------------
# Image Search
# ---------------------------------------------------------------------------

def web_search_images(query: str, max_results: int = 5) -> list:
    """Search for images via Brave Image Search. Returns list of URLs."""
    try:
        if not _get_brave_key():
            return []

        print(f"[SEARCH] Image search: {query}")
        data = _brave_get(_BRAVE_IMAGE_URL, {"q": query, "count": max_results})
        if not data:
            return []

        results = data.get("results", [])
        return [r.get("properties", {}).get("url", r.get("url", ""))
                for r in results[:max_results] if r.get("properties", {}).get("url") or r.get("url")]

    except Exception as e:
        print(f"[SEARCH] web_search_images error: {e}")
        return []


# ---------------------------------------------------------------------------
# News Search
# ---------------------------------------------------------------------------

def news_search(topic: str, groq_keys: list) -> str:
    """Search recent news on a topic via Brave News API and summarize."""
    try:
        if not _get_brave_key():
            return _NO_KEY_MSG

        print(f"[SEARCH] News search: {topic}")
        data = _brave_get(_BRAVE_NEWS_URL, {"q": topic, "count": 5})
        if not data:
            return _NO_KEY_MSG

        results = data.get("results", [])
        if not results:
            return f"No recent news about '{topic}'."

        # Build news context
        news_parts = []
        for i, r in enumerate(results[:5], 1):
            title = r.get("title", "")
            desc = r.get("description", "")
            source = r.get("meta_url", {}).get("hostname", "")
            age = r.get("age", "")
            news_parts.append(f"{i}. [{source}] {title} ({age}): {desc}")

        context = "\n".join(news_parts)

        summary = _groq_complete(
            groq_keys,
            system="Summarize these news headlines into a quick spoken briefing. "
                   "Be concise (3-5 sentences), conversational, like a news anchor.",
            user=f"News about '{topic}':\n{context}",
            max_tokens=250,
        )

        return summary or f"Latest news on {topic}: " + ". ".join(
            r.get("title", "") for r in results[:3])

    except Exception as e:
        print(f"[SEARCH] news_search error: {e}")
        return f"Couldn't search news: {e}"


# ---------------------------------------------------------------------------
# Answer Question (smart routing)
# ---------------------------------------------------------------------------

def answer_question(question: str, groq_keys: list) -> str:
    """Answer a factual question — uses web search if needed, Groq direct if simple."""
    try:
        # First try: ask Groq if this needs a web search
        needs_search = _groq_complete(
            groq_keys,
            system=(
                "Decide if this question needs a web search to answer accurately. "
                "Reply with ONLY 'search' or 'direct'. "
                "'search' = needs current/factual data (prices, news, scores, recent events, specific people). "
                "'direct' = general knowledge, math, definitions, opinions, creative tasks."
            ),
            user=question,
            max_tokens=10,
            temperature=0.0,
        )

        if needs_search and "search" in needs_search.lower():
            print(f"[SEARCH] Question needs web search: {question}")
            return web_search(question, groq_keys)

        # Direct answer from Groq
        print(f"[SEARCH] Answering directly: {question}")
        answer = _groq_complete(
            groq_keys,
            system="Answer this question concisely and conversationally (2-3 sentences). "
                   "Speak naturally as if telling a friend.",
            user=question,
            max_tokens=200,
        )
        return answer or "I'm not sure about that."

    except Exception as e:
        print(f"[SEARCH] answer_question error: {e}")
        return f"Couldn't answer that: {e}"
